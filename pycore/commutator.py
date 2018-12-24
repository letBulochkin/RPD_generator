from docx import Document
from . import struct, generator

class rpd(object):

	def __init__(self, source, rule, sample):

		self.STUDY_PLAN = struct.study_plan(source, rule)
		self.DOC = Document(sample)
		self.PREFIX = 'РАБОЧАЯ_ПРОГРАММА_'
		self.ADDED_TABLES = 0

		#self.__get_paragraphs__()
		self.TEXT = generator.get_marked_paragraphs(self.DOC)

		self.__check_file__()

	def __get_paragraphs__(self):

		self.TEXT = [[x] for x in self.DOC.paragraphs if x.text != '']  # получить все не пустые параграфы
		for i in range(len(self.TEXT)):  # добавляем к параграфам каретки с маркерами
			self.TEXT[i].append([k for k in self.TEXT[i][0].runs if "<>" in k.text])

	def create_discipline(self, index):

		self.DISCIPLINE = struct.discipline(self.STUDY_PLAN, index)

	def __write_file__(self, path, title):

		self.DOC.save(path + self.PREFIX + title + '.docx')

	def __check_file__(self):

		#[4, 1, 1]
		table_check_list = [
			[2, 0, 0], [2, 2, 6], [2, 4, 2], [2, 6, 3], [2, 8, 4], [2, 10, 5], [2, 12, 1], [4, 1, 1], [6, 1, 0],
			[6, 1, 1], [6, 1, 2], [7, 3, 0], [7, 3, 1], [7, 3, 2], [7, 3, 3], [7, 3, 4], [7, 3, 5], [7, 3, 6],
			[7, 3, 7], [8, 1, 0], [8, 1, 1], [8, 1, 2], [9, 1, 0], [9, 1, 1], [9, 1, 2], [10, 1, 0], [10, 1, 1],
			[10, 1, 2], [12, 2, 4], [13, 1, 0], [14, 1, 0], [15, 1, 0], [16, 1, 0], [16, 1, 1], [17, 1, 0],
			[17, 1, 1], [17, 1, 2], [18, 1, 0], [18, 1, 1], [18, 1, 2]
		]

		par_check_list = [
			[4, 4], [6, 4], [7, 2], [20, 1], [22, 1], [40, 2], [44, 1], [46, 3], [48, 4], [49, 2], [50, 1], [53, 1], [72, 2]
		]

		par_in_cells_check_list = [[11, 0, 0, 3], [19, 0, 0, 1]]

		for i in table_check_list:
			assert generator.check_marked_cell(
				self.DOC.tables[i[0]], i[1], i[2]), "Ошибка чтения маркера в таблице {} в клетке {}, {}".format(i[0], i[1], i[2])

		for i in par_check_list:
			assert generator.check_marked_paragraph(
				self.TEXT[i[0]][0], i[1]), "Ошибка чтения маркера в абзаце {}:".format(i[0]) + self.TEXT[i[0]][0].text[0:30] + "..."

		'''
		for i in par_in_cells_check_list:
			assert generator.check_marked_paragraph(
				self.DOC.tables[i[0]].cell(i[1], i[2]), i[3]), "Ошибка чтения маркеров в таблице {} в ячейке {}, {}".format(i[0], i[1], i[2])
		'''
		
	def produce(self, path):
		
		generator.write_to_par(self.TEXT[4], 0, self.DISCIPLINE.NAME)

		queue_string = ''
		for i in self.DISCIPLINE.COMPETENCIES:
			queue_string += ' ' + i['code'].strip() + ' ' + '"'.strip() + i['descrp'].strip() + '"'.strip()
			if str(i['part']) != "" and str(i['part']) != None:
				queue_string += ' в части ' + i['part'].strip()  # должно быть более элегантное решение
			queue_string += ', '

		generator.write_to_par(self.TEXT[4], 1, queue_string)
		
		generator.write_to_par(self.TEXT[4], 2, self.STUDY_PLAN.FIELD_OF_KNOW)
		generator.write_to_par(self.TEXT[4], 3, self.STUDY_PLAN.PROFILE)

		generator.write_to_par(self.TEXT[6], 0, self.DISCIPLINE.NAME)
		if not self.DISCIPLINE.OBLIGATION:
			generator.write_to_par(self.TEXT[6], 1, 'по выбору')
		else:
			generator.write_to_par(self.TEXT[6], 1, '')
		if self.DISCIPLINE.PART:
			generator.write_to_par(self.TEXT[6], 2, 'базовой части')
		else:
			generator.write_to_par(self.TEXT[6], 2, 'вариативной части')

		generator.write_to_par(self.TEXT[6], 3, self.STUDY_PLAN.FIELD_OF_KNOW)

		sum_zach, sum_total = 0, 0
		for i in self.DISCIPLINE.STUDY_HOURS:
			sum_zach += i[1]['zach_ed']
			sum_total += i[1]['total']
		generator.write_to_par(self.TEXT[7], 0, str(sum_zach))
		generator.write_to_par(self.TEXT[7], 1, str(sum_total))

		generator.write_to_par(self.TEXT[20], 0, self.DISCIPLINE.NAME)
		generator.write_to_par(self.TEXT[22], 0, self.DISCIPLINE.NAME)

		generator.write_to_par(self.TEXT[40], 0, self.DISCIPLINE.NAME)
		generator.write_to_par(self.TEXT[40], 1, self.STUDY_PLAN.FIELD_OF_KNOW)

		generator.write_to_par(self.TEXT[44], 0, self.DISCIPLINE.NAME)

		generator.write_to_par(self.TEXT[46], 0, self.DISCIPLINE.NAME)
		generator.write_to_par(self.TEXT[46], 1, ', '.join(i['code'] for i in self.DISCIPLINE.COMPETENCIES))
		generator.write_to_par(self.TEXT[46], 2, self.STUDY_PLAN.FIELD_OF_KNOW)

		generator.write_to_par(self.TEXT[48], 0, self.DISCIPLINE.NAME)
		if not self.DISCIPLINE.OBLIGATION:
			generator.write_to_par(self.TEXT[48], 1, 'по выбору')
		else:
			generator.write_to_par(self.TEXT[48], 1, '')
		if self.DISCIPLINE.PART:
			generator.write_to_par(self.TEXT[48], 2, 'базовой части')
		else:
			generator.write_to_par(self.TEXT[48], 2, 'вариативной части')

		generator.write_to_par(self.TEXT[48], 3, self.STUDY_PLAN.FIELD_OF_KNOW)

		generator.write_to_par(self.TEXT[49], 0, str(sum_zach))
		generator.write_to_par(self.TEXT[49], 1, str(sum_total))

		queue_string = ''
		for i in self.DISCIPLINE.SEMESTERS:
			if i[3] == True:
				queue_string += "зачет"
				break
		for i in self.DISCIPLINE.SEMESTERS:
			if i[2] == True:
				queue_string += ", экзамен"
				break
		generator.write_to_par(self.TEXT[50], 0, queue_string)

		generator.write_to_par(self.TEXT[53], 0, self.DISCIPLINE.NAME)

		generator.write_to_par(self.TEXT[72], 0, self.DISCIPLINE.NAME)
		generator.write_to_par(self.TEXT[72], 1, queue_string)

		generator.write_to_cell(self.DOC.tables[2], 0, 0, self.DISCIPLINE.NAME)
		generator.write_to_cell(self.DOC.tables[2], 2, 6, self.DISCIPLINE.STUDY_PLAN.FIELD_OF_KNOW)
		generator.write_to_cell(self.DOC.tables[2], 4, 2, self.DISCIPLINE.STUDY_PLAN.PROFILE)
		generator.write_to_cell(self.DOC.tables[2], 6, 3, self.DISCIPLINE.STUDY_PLAN.INSTITUTE)
		generator.write_to_cell(self.DOC.tables[2], 8, 4, self.DISCIPLINE.STUDY_PLAN.EDU_FORMAT)
		generator.write_to_cell(self.DOC.tables[2], 10, 5, self.DISCIPLINE.STUDY_PLAN.EDU_PROG)
		generator.write_to_cell(self.DOC.tables[2], 12, 1, self.DISCIPLINE.STUDY_PLAN.CATHEDRA)

		generator.write_to_cell(self.DOC.tables[4], 1, 1, self.DISCIPLINE.STUDY_PLAN.CATHEDRA)

		queue_list = []
		for i in self.DISCIPLINE.COMPETENCIES:
			q = []
			q.append(i['code'])
			s = '"'.strip() + i['descrp'].strip() + '"'.strip()
			if str(i['part']) != "":
				s += ' в части ' + i['part'].strip()
			q.append(s)
			s = ''
			if str(i['to_know']) != "" and str(i['to_know']) != None: 
				s += 'Знать ' + i['to_know'].strip() + '\n'
			if str(i['to_can']) != "" and str(i['to_can']) != None:
				s += 'Уметь ' + i['to_can'].strip() + '\n'
			if str(i['to_be_able']) != "" and str(i['to_be_able']) != None: 
				s += 'Владеть ' + i['to_be_able'].strip()
			q.append(s)
			queue_list.append(q)
		generator.seq_write_to_table(self.DOC.tables[6], queue_list)

		# generator.seq_write_to_table(self.DOC.tables[6], self.DISCIPLINE.COMPETENCIES)

		for i in range(len(self.DISCIPLINE.STUDY_HOURS)):
			print('opa!')
			t = self.DOC.tables[7]
			base_par = [i for i in self.DOC.paragraphs if '4.2' in i.text]
			parg = base_par[0].insert_paragraph_before('4.1.' + str(i+1) + 
				' Семестр ' + str(self.DISCIPLINE.STUDY_HOURS[i][0]))
			generator.copy_table_after(t, parg)
			self.ADDED_TABLES += 1
			t = self.DOC.tables[7 + i + 1]
			queue_list = []
			for k in self.DISCIPLINE.SEMESTERS:
				if k[0] == self.DISCIPLINE.STUDY_HOURS[i][0]:
					for l in k[1].MODULES:
						queue_list.append([
							str(l['num']),
							str(l['lect'] + l['lab'] + l['pract'] + l['sam']),
							str(l['lect']),
							str(l['lab']),
							str(l['pract']),
							str(l['sam']),
							'',
							''
						])
			
			if self.DISCIPLINE.SEMESTERS[i][1].EXAM:
				control_str = 'Экзамен'
			elif self.DISCIPLINE.SEMESTERS[i][1].ZACHET:
				control_str = 'Зачет'
			else:
				control_str = ''
			
			queue_list.append([
				'Всего: ', 
				str(self.DISCIPLINE.STUDY_HOURS[i][1]['total']), 
				str(self.DISCIPLINE.STUDY_HOURS[i][1]['lect']),
				str(self.DISCIPLINE.STUDY_HOURS[i][1]['lab']),
				str(self.DISCIPLINE.STUDY_HOURS[i][1]['pract']),
				str(self.DISCIPLINE.STUDY_HOURS[i][1]['sam']),
				str(self.DISCIPLINE.STUDY_HOURS[i][1]['krpa'] + self.DISCIPLINE.STUDY_HOURS[i][1]['control']),
				control_str
			])
			generator.seq_write_to_table(t, queue_list)
		generator.remove_table(self.DOC.tables[7])
		self.ADDED_TABLES -= 1
		
		sdv = self.ADDED_TABLES
		for i in range(len(self.DISCIPLINE.SEMESTERS)):
			t = self.DOC.tables[8 + sdv]
			base_par = [i for i in self.DOC.paragraphs if '4.3' in i.text]
			parg = base_par[0].insert_paragraph_before('4.2.' + str(i+1) + 
				' Семестр ' + str(self.DISCIPLINE.SEMESTERS[i][0]))
			generator.copy_table_after(t, parg)
			self.ADDED_TABLES += 1
			t = self.DOC.tables[8 + sdv + i + 1]
			queue_list = []
			for k in self.DISCIPLINE.SEMESTERS[i][1].MODULES:
				queue_list.append([str(k['num']), str(k['descr'])])
			generator.seq_write_to_table(t, queue_list)
		generator.remove_table(self.DOC.tables[8 + sdv])
		self.ADDED_TABLES -= 1

		if self.DISCIPLINE.LABS != []:
			generator.seq_write_to_table(self.DOC.tables[9 + self.ADDED_TABLES], self.DISCIPLINE.LABS)
		else:
			generator.remove_table(self.DOC.tables[9 + self.ADDED_TABLES])
			self.ADDED_TABLES -= 1

		if self.DISCIPLINE.PRACT != []:
			generator.seq_write_to_table(self.DOC.tables[10 + self.ADDED_TABLES], self.DISCIPLINE.PRACT)
		else:
			generator.remove_table(self.DOC.tables[10 + self.ADDED_TABLES])
			self.ADDED_TABLES -= 1

		reserve_t = self.DOC.tables[11 + self.ADDED_TABLES]
		base_par = [i for i in self.DOC.paragraphs if 'Комплекты контрольных заданий' in i.text]
		queue_string = ''
		for k in self.DISCIPLINE.COMPETENCIES:
			queue_string += ' ' + k['code'].strip()
		for i in self.DISCIPLINE.SEMESTERS:
			if i[3] == True:
				parg = base_par[0].insert_paragraph_before(' ')
				generator.copy_table_after(reserve_t, parg)
				self.ADDED_TABLES += 1
				t = self.DOC.tables[11 + self.ADDED_TABLES]
				pars = generator.get_marked_paragraphs(t.cell(0,0))
				generator.write_to_par(pars[0], 0, 'зачету')
				generator.write_to_par(pars[0], 1, str(i[0]))
				generator.write_to_par(pars[0], 2, queue_string)
				break
		for i in self.DISCIPLINE.SEMESTERS:
			if i[2] == True:
				parg = base_par[0].insert_paragraph_before(' ')
				generator.copy_table_after(reserve_t, parg)
				self.ADDED_TABLES += 1
				t = self.DOC.tables[11 + self.ADDED_TABLES]
				pars = generator.get_marked_paragraphs(t.cell(0,0))
				generator.write_to_par(pars[0], 0, 'экзамену')
				generator.write_to_par(pars[0], 1, str(i[0]))
				generator.write_to_par(pars[0], 2, queue_string)
				break
		generator.remove_table(reserve_t)
		self.ADDED_TABLES -= 1

		queue_string = ''
		for i in self.DISCIPLINE.SEMESTERS:
			if i[3] == True:
				queue_string += "зачет"
				break
		for i in self.DISCIPLINE.SEMESTERS:
			if i[2] == True:
				queue_string += ", экзамен"
				break
		generator.write_to_cell(self.DOC.tables[12 + self.ADDED_TABLES], 2, 4, queue_string)

		queue_string = ''
		for i in self.DISCIPLINE.COMPETENCIES:
			if str(i['to_know']) != "" and str(i['to_know']) != None:
				queue_string += ' ' + str(i['to_know']) + ';'
		generator.write_to_cell(self.DOC.tables[13 + self.ADDED_TABLES], 1, 0, queue_string)

		queue_string = ''
		for i in self.DISCIPLINE.COMPETENCIES:
			if str(i['to_can']) != "" and str(i['to_can']) != None:
				queue_string += ' ' + str(i['to_can']) + ';'
		generator.write_to_cell(self.DOC.tables[14 + self.ADDED_TABLES], 1, 0, queue_string)

		queue_string = ''
		for i in self.DISCIPLINE.COMPETENCIES:
			if str(i['to_be_able']) != "" and str(i['to_be_able']) != None:
				queue_string += ' ' + str(i['to_be_able']) + ';'
		generator.write_to_cell(self.DOC.tables[15 + self.ADDED_TABLES], 1, 0, queue_string)

		generator.seq_write_to_table(self.DOC.tables[16 + self.ADDED_TABLES], self.STUDY_PLAN.PRACTICE_TYPES)

		queue_list = []
		for i in self.DISCIPLINE.COMPETENCIES:
			q = []
			q.append(i['code'])
			s = '"'.strip() + i['descrp'].strip() + '"'.strip()
			if str(i['part']) != "":
				s += ' в части ' + i['part'].strip()
			q.append(s)
			s = ''
			if str(i['to_know']) != "" and str(i['to_know']) != None: 
				s += 'Знать ' + i['to_know'].strip() + '\n'
			if str(i['to_can']) != "" and str(i['to_can']) != None:
				s += 'Уметь ' + i['to_can'].strip() + '\n'
			if str(i['to_be_able']) != "" and str(i['to_be_able']) != None: 
				s += 'Владеть ' + i['to_be_able'].strip()
			q.append(s)
			queue_list.append(q)
		if queue_list != []:
			generator.seq_write_to_table(self.DOC.tables[17 + self.ADDED_TABLES], queue_list)

		queue_list = []
		for i in self.DISCIPLINE.COMPETENCIES:
			if str(i['to_know']) != "" and str(i['to_know']) != None:
				q = []
				q.append('Знать (' + i['code'] + ')')
				q.append('Знать ' + i['to_know'])
				q.append('Выполнение устных/письменных заданий')
				queue_list.append(q)
			if str(i['to_can']) != "" and str(i['to_can']) != None:
				q = []
				q.append('Уметь (' + i['code'] + ')')
				q.append('Уметь ' + i['to_can'])
				q.append('Выполнение устных/письменных заданий')
				queue_list.append(q)
			if str(i['to_be_able']) != "" and str(i['to_be_able']) != None:
				q = []
				q.append('Владеть (' + i['code'] + ')')
				q.append('Владеть ' + i['to_be_able'])
				q.append('Выполнение устных/письменных заданий')
				queue_list.append(q)
		if queue_list != []:
			generator.seq_write_to_table(self.DOC.tables[18 + self.ADDED_TABLES], queue_list)

		reserve_t = self.DOC.tables[19 + self.ADDED_TABLES]
		base_par = [i for i in self.DOC.paragraphs if 'Шкала оценивания:' in i.text]
		for i in self.DISCIPLINE.SEMESTERS:
			if i[3] == True:
				parg = base_par[0].insert_paragraph_before(' ')
				generator.copy_table_after(reserve_t, parg)
				self.ADDED_TABLES += 1
				t = self.DOC.tables[19 + self.ADDED_TABLES]
				pars = generator.get_marked_paragraphs(t.cell(0,0))
				generator.write_to_par(pars[0], 0, 'зачету')
				break
		for i in self.DISCIPLINE.SEMESTERS:
			if i[2] == True:
				parg = base_par[0].insert_paragraph_before(' ')
				generator.copy_table_after(reserve_t, parg)
				self.ADDED_TABLES += 1
				t = self.DOC.tables[19 + self.ADDED_TABLES]
				pars = generator.get_marked_paragraphs(t.cell(0,0))
				generator.write_to_par(pars[0], 0, 'экзамену')
				break
		generator.remove_table(reserve_t)
		self.ADDED_TABLES -= 1

		#self.__write_file__("C:\\Users\\Anton Firsov\\Documents\\Python\\RPD_generator\\data\\", self.DISCIPLINE.INDEX)
		self.__write_file__(path, self.DISCIPLINE.INDEX + '_' + self.DISCIPLINE.NAME)
