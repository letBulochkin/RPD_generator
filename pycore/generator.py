from docx import Document
from docx.shared import Pt
from copy import deepcopy

def check_marked_cell(table, row, col):
	"""Checks whether marker in given table cell exists.

	Args:
		table (...): docx table object
		row (int): table row
		col (int): table column

	Returns:
		bool: True if marker exists, False if not
	"""

	if table.cell(row, col).paragraphs[0].runs[0].text == "<>":  # если в указанной ячейке только маркер
		return True
	else: 
		return False

def check_marked_paragraph(paragraph, number):
	"""Checks if there are a number of markers in given paragraph.

	Args:
		paragraph (...): docx paragraph object
		number (int): quantity of markers

	Returns:
		bool: True if all markers exist and quantity matches with given
	"""

	q = 0  # счетчик найденных маркеров
	chars = '<> '  # возможные символы в каретке

	for i in range(len(paragraph.runs)):
		if "<>" in paragraph.runs[i].text:  # если в тексте каретки встречается маркер
			for c in paragraph.runs[i].text:  # проверяем каждый символ в каретке
				if c not in chars:  # если он не входит в список разрешенных символов
					return False
			q += 1  # если проверка пройдена, увеличиваем счетчик
		elif "<" in paragraph.runs[i].text and ">" in paragraph.runs[i+1].text:  # если маркер разделен на две соседние каретки
			for c in paragraph.runs[i].text:  # проверяем каждую из кареток
				if c not in chars:
					return False
			for c in paragraph.runs[i+1].text:
				if c not in chars:
					return False
			q += 1

	if q != number:  # если количество маркеров не совпало с указанным в выводе
		return False
	else:
		return True

def get_marked_cells(table):
	"""Returns docx table cells that marked with special marker.

	Args:
		table (...): docx table object

	Returns:
		list: list of lists containing marked cell coordinate [row, column]
	"""

	res = []

	for i in range(len(table.rows)):
		for k in range(len(table.row_cells(i))):
			if table.cell(i, k).text == '<>':  # marker
				res.append([i, k])
				break  # но зачем??

	return res

def get_marked_paragraphs(doc):
	"""Returns list of paragraphs with special marker.

	Args:
		doc (...): docx Document object

	Returns:
		list: list of lists containing paragraph object and list of marked runs
	"""

	res = [[x] for x in doc.paragraphs if x.text != '']  # получаем все непустые параграфы

	for i in range(len(res)):
		q = []  # подготавливаем список маркеров
		for k in range(len(res[i][0].runs)):
			if "<>" in res[i][0].runs[k].text:  # если в тексте каретки встречается маркер
				q.append(res[i][0].runs[k])
			elif "<" in res[i][0].runs[k].text and ">" in res[i][0].runs[k+1].text:  # сли маркер разделен на две сосендние каретки
				res[i][0].runs[k+1].clear()  # удаляем содержимое второй каретки
				q.append(res[i][0].runs[k])  # и сохраняем в итоговый список первую 
		if q != []:  # если найдены маркеры
			res[i].append(q)

	return res

def insert_table_after(table, paragraph):
	"""Inserts table after paragraph.

	Credits: https://github.com/python-openxml/python-docx/issues/156
	ATTENTION: This method modifies list of paragraphs. 

	Args:
		table (...): docx table to copy
		paragraph (...): docx paragraph to paste table after
	"""

	tbl, p = table._tbl, paragraph._p
	p.addnext(tbl)

def copy_table_after(table, paragraph):
	"""Copies table after paragraph.

	Credits: https://github.com/python-openxml/python-docx/issues/156
	ATTENTION: This method modifies list of paragraphs. 

	Args:
		table (...): docx table to copy
		paragraph (...): docx paragraph to paste table after
	"""

	tbl, p = table._tbl, paragraph._p
	new_tbl = deepcopy(tbl)
	p.addnext(new_tbl)

def remove_table(table):
	"""Removes table. Credit: https://stackoverflow.com/questions/47716792/python-docx-delete-table-from-document#comment82455626_47743738"""

	table._element.getparent().remove(table._element)

def seq_write_to_table(table, data):
	"""Sequential writing to self-created table.
	
	Последовательная запись данных вложенного списка в таблицу с добавлением новых ячеек. Метод опеределяет 
	маркированные в таблице ячейки, затем последовательно вставляет в маркированные ячейки соответственно 
	идущие элементы списка. Под каждый новый список создается новая строка в таблице.

	Args:
		table (...): table to overwrite
		data (list): list of lists of the same size
	"""

	markers = []  # список маркеров в заготовке таблицы

	for i in range(len(table.rows)):
		for k in range(len(table.row_cells(i))):
			if table.cell(i, k).text == '<>':  # marker
				markers.append([i, k, table.cell(i, k).paragraphs[0].runs[0].style])  # записываем координаты клетки с маркером
				table.cell(i, k).text = ''  #и стиль каретки. Ячейку с маркером обнуляем

	for i in range(len(data[0])):  # идем сначала по списку вглубь
		for k in range(len(data)):  # потом вширь
			try:
				table.cell(markers[i][0] + k, markers[i][1])  #ячейка существует?
			except IndexError:
				row = table.add_row()  # если нет, добавляем ряд
				[i.add_paragraph() for i in row.cells]  # вставляем в каждую ячейку нового ряда по параграфу
				# добавляем в параграф каретку с текстом и применяем стиль
				table.cell(markers[i][0] + k, markers[i][1]).paragraphs[0].add_run(data[k][i], markers[i][2])
			else:  # если существует, то просто добавляем текст
				table.cell(markers[i][0] + k, markers[i][1]).paragraphs[0].add_run(data[k][i], markers[i][2])

def write_to_cell(table, row, col, text):
	"""Writing text to cell of the table"""

	table.cell(row, col).paragraphs[0].runs[0].text = text

def write_to_par(par, pos, text):
	"""Overwrites marked run in a paragraph.

	TODO: This approach has a major flaw!

	Args:
		par (list): list containing paragraph object and list of marked runs,
			included in it
		pos (int): position of marker to overwrite, 0-based
		text (str): text to write to marker 
	"""

	par[1][pos].text = text
